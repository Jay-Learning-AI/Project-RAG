from pathlib import Path
from typing import Iterator

from docx import Document as WordDocument
from docx.document import Document as WordProcessingDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredHTMLLoader
)

SUPPORTED_EXTENSIONS = {
    ".pdf": PyPDFLoader,
    ".docx": None,
    ".txt": TextLoader,
    ".html": UnstructuredHTMLLoader,
}

DRAWING_NAMESPACES = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def _iter_block_items(document: WordProcessingDocument) -> Iterator[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _extract_image_names_from_element(document: WordProcessingDocument, element) -> list[str]:
    image_names = []
    for blip in element.xpath(".//a:blip", namespaces=DRAWING_NAMESPACES):
        rel_id = blip.get(qn("r:embed"))
        if not rel_id or rel_id not in document.part.rels:
            continue

        target_ref = document.part.rels[rel_id].target_ref
        image_names.append(Path(target_ref).name)

    # Preserve order while removing duplicates.
    return list(dict.fromkeys(image_names))


def _paragraph_to_block(document: WordProcessingDocument, paragraph: Paragraph) -> dict:
    return {
        "text": paragraph.text.strip(),
        "image_names": _extract_image_names_from_element(document, paragraph._element),
    }


def _table_to_block(document: WordProcessingDocument, table: Table) -> dict:
    rows = []
    for row in table.rows:
        values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if values:
            rows.append(" | ".join(values))

    return {
        "text": "\n".join(rows).strip(),
        "image_names": _extract_image_names_from_element(document, table._element),
    }


def _load_docx_documents(file_path: Path) -> list[Document]:
    document = WordDocument(str(file_path))
    blocks = []

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            blocks.append(_paragraph_to_block(document, block))
        else:
            blocks.append(_table_to_block(document, block))

    docs = []
    for index, block in enumerate(blocks):
        text = block["text"]
        if not text:
            continue

        image_names = list(block["image_names"])
        if index > 0 and not blocks[index - 1]["text"]:
            image_names.extend(blocks[index - 1]["image_names"])
        if index + 1 < len(blocks) and not blocks[index + 1]["text"]:
            image_names.extend(blocks[index + 1]["image_names"])

        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": file_path.name,
                    "paragraph_index": index,
                    "image_names": list(dict.fromkeys(image_names)),
                },
            )
        )

    if docs:
        return docs

    return [
        Document(
            page_content=document.paragraphs[0].text.strip() if document.paragraphs else "",
            metadata={"source": file_path.name, "paragraph_index": 0, "image_names": []},
        )
    ]


def load_documents(file_path: str):
    """
    Load documents of any supported type and return
    LangChain Document objects.
    """
    file_path = Path(file_path)
    extension = file_path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {extension}. "
            f"Supported types: {list(SUPPORTED_EXTENSIONS.keys())}"
        )

    if extension == ".docx":
        return _load_docx_documents(file_path)

    loader_class = SUPPORTED_EXTENSIONS[extension]

    # Special handling for TextLoader
    if extension == ".txt":
        loader = loader_class(str(file_path), encoding="utf-8")
    else:
        loader = loader_class(str(file_path))

    return loader.load()


if __name__ == "__main__":
    import os

    DATA_DIR = os.path.join(os.path.dirname(__file__), "..", "data", "documents")
    files = os.listdir(DATA_DIR)

    if not files:
        print("No files found in data/documents/")
    else:
        for filename in files:
            file_path = os.path.join(DATA_DIR, filename)
            print(f"\n{'='*50}")
            print(f"Loading: {filename}")
            try:
                docs = load_documents(file_path)
                print(f"✅ Loaded {len(docs)} document chunk(s)")
                for i, doc in enumerate(docs[:2]):  # preview first 2 chunks
                    preview = doc.page_content[:300].replace("\n", " ").strip()
                    print(f"  [Chunk {i+1}] {preview}...")
            except ValueError as e:
                print(f"⚠️  Skipped: {e}")
            except Exception as e:
                print(f"❌ Error loading {filename}: {e}")
